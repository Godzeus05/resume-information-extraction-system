"""Document text extraction for PDF and DOCX resumes."""

from pathlib import Path


def extract_text(file_path: str) -> str:
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    if path.suffix.lower() == ".pdf":
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    if path.suffix.lower() == ".docx":
        from docx import Document
        doc = Document(str(path))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(parts)

    raise ValueError("Unsupported file type. Only PDF and DOCX are supported.")
